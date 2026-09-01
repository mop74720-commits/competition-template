#!/usr/bin/env python3
from __future__ import annotations
import argparse, csv, hashlib, re, shutil, subprocess, tempfile, zipfile
from pathlib import Path
import xml.etree.ElementTree as ET

ROOT=Path(__file__).resolve().parents[1]
WNS='{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
RANGE_RE=re.compile(r'(?i)(?:不少于|不低于|至少|不超过|不高于|至多|最多|上限|下限|范围|介于|[<>≤≥])|(?:\d[\d,.]*(?:\s*)(?:-|~|～|—|–|至|到)(?:\s*)\d[\d,.]*)|(?:\d[\d,.]*\s*(?:%|％|万元|亿元|元|年|月|日|小时|分钟|秒|kg|g|km|m|人|家|次|台|件|吨|L|ml))')

def sha256(p):
    h=hashlib.sha256()
    with Path(p).open('rb') as f:
        for b in iter(lambda:f.read(1024*1024),b''): h.update(b)
    return h.hexdigest()

def unique_dest(dst_dir,src):
    dst=dst_dir/src.name
    if not dst.exists(): return dst
    if sha256(dst)==sha256(src): return dst
    i=2
    while True:
        cand=dst_dir/f'{src.stem}-{i}{src.suffix}'
        if not cand.exists(): return cand
        i+=1

def docx_text(path):
    py_text=''; xml_text=''
    try:
        from docx import Document
        d=Document(path); parts=[]
        parts += [p.text for p in d.paragraphs]
        for table in d.tables:
            for row in table.rows: parts.append('\t'.join(c.text for c in row.cells))
        py_text='\n'.join(parts)
    except Exception: pass
    try:
        with zipfile.ZipFile(path) as z:
            parts=[]
            for name in z.namelist():
                if name.startswith('word/') and name.endswith('.xml'):
                    try:
                        root=ET.fromstring(z.read(name)); parts.extend((x.text or '') for x in root.iter(WNS+'t'))
                    except Exception: pass
            xml_text='\n'.join(parts)
    except Exception: pass
    best=xml_text if len(xml_text)>len(py_text) else py_text
    return best, {'python_docx_chars':len(py_text),'ooxml_all_text_chars':len(xml_text),'possible_object_text':len(xml_text)>len(py_text)+10}

def pdf_text(path,render_dir=None):
    try:
        import fitz
    except Exception as e: return '', {'error':f'PyMuPDF unavailable: {e}'}
    d=fitz.open(path); chunks=[]; pages=[]
    for i,p in enumerate(d):
        txt=p.get_text('text') or ''; chunks.append(txt); pages.append(len(txt.strip()))
        if render_dir:
            render_dir.mkdir(parents=True,exist_ok=True); pix=p.get_pixmap(matrix=fitz.Matrix(1.5,1.5),alpha=False); pix.save(render_dir/f'page-{i+1}.png')
    d.close(); return '\n'.join(chunks), {'pages':len(pages),'page_text_chars':pages,'low_text_pages':[i+1 for i,n in enumerate(pages) if n<40]}

def convert_docx_pdf(path,outdir):
    exe=shutil.which('libreoffice') or shutil.which('soffice')
    if not exe:return None,'LibreOffice unavailable'
    r=subprocess.run([exe,'--headless','--convert-to','pdf','--outdir',str(outdir),str(path)],stdout=subprocess.PIPE,stderr=subprocess.STDOUT,text=True,timeout=90)
    p=outdir/(path.stem+'.pdf')
    return (p if r.returncode==0 and p.exists() else None),r.stdout[-1200:]

def xlsx_meta(path):
    try:
        with zipfile.ZipFile(path) as z:
            wb=ET.fromstring(z.read('xl/workbook.xml')); ns='{http://schemas.openxmlformats.org/spreadsheetml/2006/main}'
            names=[s.attrib.get('name','') for s in wb.find(ns+'sheets')]
            return {'sheets':names}
    except Exception as e:return {'error':str(e)}

def candidates(text):
    out=[]
    for raw in text.splitlines():
        line=' '.join(raw.split())
        if line and any(ch.isdigit() for ch in line) and RANGE_RE.search(line):
            if line not in out: out.append(line[:500])
    return out[:300]

def main():
    ap=argparse.ArgumentParser(description='Import contest statement/data and generate advisory hard-constraint preflight. No OCR and no automatic FACT promotion.')
    ap.add_argument('--statement',action='append',default=[]); ap.add_argument('--data',action='append',default=[]); ap.add_argument('--render',action='store_true'); ap.add_argument('--no-copy',action='store_true')
    a=ap.parse_args(); manifest=[]; report=[]; all_candidates=[]
    stmt_dir=ROOT/'problem/statement'; data_dir=ROOT/'data/raw'; stmt_dir.mkdir(parents=True,exist_ok=True); data_dir.mkdir(parents=True,exist_ok=True)
    render_root=ROOT/'problem/_preflight_render'
    for role,items,dstdir in [('statement',a.statement,stmt_dir),('data',a.data,data_dir)]:
        for raw in items:
            src=Path(raw).resolve()
            if not src.exists(): raise SystemExit(f'missing source: {src}')
            dst=src if a.no_copy else unique_dest(dstdir,src)
            if not a.no_copy and (not dst.exists() or sha256(dst)!=sha256(src)): shutil.copy2(src,dst)
            rel=str(dst.relative_to(ROOT)) if dst.is_relative_to(ROOT) else str(dst)
            ext=dst.suffix.lower(); extraction='metadata-only'; notes=[]; text=''; meta={}
            if role=='statement':
                if ext=='.pdf': text,meta=pdf_text(dst,render_root/dst.stem if a.render else None); extraction='pdf-text'
                elif ext=='.docx':
                    text,meta=docx_text(dst); extraction='docx-ooxml'
                    if a.render:
                        with tempfile.TemporaryDirectory(prefix='contest_docx_') as td:
                            pdf,log=convert_docx_pdf(dst,Path(td))
                            if pdf:
                                ptxt,pmeta=pdf_text(pdf,render_root/dst.stem); meta['rendered_pdf_chars']=len(ptxt); meta['rendered_pdf_pages']=pmeta.get('pages'); meta['render_low_text_pages']=pmeta.get('low_text_pages');
                                if len(ptxt)>len(text)*1.05+20: notes.append('rendered PDF contains materially more text than DOCX extraction; inspect page images')
                            else: notes.append('DOCX render unavailable/failed: '+log.replace('\n',' ')[:300])
                elif ext in {'.txt','.md'}: text=dst.read_text(encoding='utf-8-sig',errors='replace'); extraction='plain-text'
                all_candidates.extend((rel,x) for x in candidates(text))
                if meta.get('possible_object_text'): notes.append('OOXML contains more text than python-docx; possible textbox/object text was recovered')
                if meta.get('low_text_pages'): notes.append('low-text PDF pages require visual review: '+str(meta['low_text_pages']))
            elif ext in {'.xlsx','.xlsm'}: meta=xlsx_meta(dst); extraction='xlsx-metadata'
            manifest.append({'role':role,'relative_path':rel,'sha256':sha256(dst),'bytes':dst.stat().st_size,'extraction':extraction,'notes':'; '.join(notes)})
            report.append((rel,extraction,meta,notes))
    mf=ROOT/'problem/SOURCE_MANIFEST.csv'
    with mf.open('w',encoding='utf-8-sig',newline='') as f:
        w=csv.DictWriter(f,fieldnames=['role','relative_path','sha256','bytes','extraction','notes']); w.writeheader(); w.writerows(manifest)
    rp=ROOT/'problem/PREFLIGHT_REPORT.md'
    lines=['# Contest Import Preflight','','> Advisory only. Candidates below are NOT formal FACTS until checked against the original statement/render.','', '## Imported sources','']
    for rel,extraction,meta,notes in report:
        lines += [f'### `{rel}`',f'- extraction: `{extraction}`',f'- metadata: `{meta}`']
        for n in notes: lines.append(f'- **REVIEW:** {n}')
        lines.append('')
    lines += ['## Numeric / range / unit candidates','']
    if all_candidates:
        for rel,line in all_candidates: lines.append(f'- `{rel}`: {line}')
    else: lines.append('- No candidates detected automatically. This does not prove the statement contains no hard numeric constraints.')
    lines += ['', '## Human preflight checklist','', '- Compare all detected ranges, inequalities, units and upper/lower bounds with the original page/render.', '- Inspect low-text pages and diagrams/tables visually; no OCR is performed by default.', '- Move only verified facts into `FACTS.md`; assumptions and interpretations remain separate.', '- Check that attachment units/keys align with the statement before modeling.','']
    rp.write_text('\n'.join(lines),encoding='utf-8')
    print(f'IMPORTED={len(manifest)} manifest={mf.relative_to(ROOT)} report={rp.relative_to(ROOT)}')
if __name__=='__main__': main()
